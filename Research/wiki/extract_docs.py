import os
import sys
import docx

sys.stdout.reconfigure(encoding='utf-8')

def extract_text_from_docx(path):
    doc = docx.Document(path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                if cell.text.strip():
                    row_text.append(cell.text.strip())
            if row_text:
                texts.append(" | ".join(row_text))
    return "\n".join(texts)

# Find files
research_dir = r"D:\laptrinhdidong\DoAn3\Research"
output_file = os.path.join(research_dir, "wiki", "docs_raw.txt")

files = [f for f in os.listdir(research_dir) if f.endswith(".docx")]
files.sort()

with open(output_file, "w", encoding="utf-8") as out:
    for i, fname in enumerate(files):
        path = os.path.join(research_dir, fname)
        print(f"Processing [{i+1}/{len(files)}]: {fname}")
        out.write("=" * 80 + "\n")
        out.write(f"FILE {i+1}: {fname}\n")
        out.write("=" * 80 + "\n")
        try:
            text = extract_text_from_docx(path)
            out.write(text)
        except Exception as e:
            out.write(f"[ERROR reading file: {e}]")
        out.write("\n\n")

print(f"Done. Output: {output_file}")
